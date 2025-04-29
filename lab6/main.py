import tkinter as tk
from tkinter import ttk, messagebox, filedialog
from docx import Document
from docx.shared import Pt
from calc.room import calculate_room
from calc.apartment import calculate_apartment
from calc.building import calculate_building

def doc_report(path, data):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(14)
    doc.add_heading("Отчёт по расчёту помещения", 2)
    for key, value in data.items():
        doc.add_paragraph(f"{key}: {value}")
    doc.save(path)

def calculate():
    try:
        length = float(length_entry.get())
        width = float(width_entry.get())
        height = float(height_entry.get())
        rooms = int(rooms_entry.get() or 1)
        floors = int(floors_entry.get() or 1)
        building_type = type_var.get()

        if building_type == "Комната":
            area, heat = calculate_room(length, width, height)
        elif building_type == "Квартира":
            area, heat = calculate_apartment(length, width, height, rooms)
        elif building_type == "Многоэтажный дом":
            area, heat = calculate_building(length, width, height, rooms, floors)
        else:
            raise ValueError("Неизвестный тип помещения")

        result_var.set(f"Площадь: {area} м²\nТепловая мощность: {heat} Вт")
    except Exception as e:
        messagebox.showerror("Ошибка", str(e))

def save_report(format):
    file = filedialog.asksaveasfilename(defaultextension=f".{format}", filetypes=[(format.upper(), f"*.{format}")])
    if not file:
        return
    length = float(length_entry.get())
    width = float(width_entry.get())
    height = float(height_entry.get())
    rooms = int(rooms_entry.get() or 1)
    floors = int(floors_entry.get() or 1)
    building_type = type_var.get()

    if building_type == "Комната":
        area, heat = calculate_room(length, width, height)
    elif building_type == "Квартира":
        area, heat = calculate_apartment(length, width, height, rooms)
    elif building_type == "Многоэтажный дом":
        area, heat = calculate_building(length, width, height, rooms, floors)

    data = {
        "Тип": building_type,
        "Длина": length,
        "Ширина": width,
        "Высота": height,
        "Комнат": rooms,
        "Этажей": floors,
        "Площадь": area,
        "Тепловая мощность": heat,
    }

    if format == "doc":
        doc_report(file, data)
        
app = tk.Tk()
app.title("Расчёт помещения")

tk.Label(app, text="Длина (м):", font='Times 14').grid(row=0, column=0)
length_entry = tk.Entry(app)
length_entry.grid(row=0, column=1)

tk.Label(app, text="Ширина (м):", font='Times 14').grid(row=1, column=0)
width_entry = tk.Entry(app)
width_entry.grid(row=1, column=1)

tk.Label(app, text="Высота (м):", font='Times 14').grid(row=2, column=0)
height_entry = tk.Entry(app)
height_entry.grid(row=2, column=1)

tk.Label(app, text="Кол-во комнат:", font='Times 14').grid(row=3, column=0)
rooms_entry = tk.Entry(app)
rooms_entry.grid(row=3, column=1)

tk.Label(app, text="Кол-во этажей:", font='Times 14').grid(row=4, column=0)
floors_entry = tk.Entry(app)
floors_entry.grid(row=4, column=1)

tk.Label(app, text="Тип помещения:", font='Times 14').grid(row=5, column=0)
type_var = tk.StringVar(value="Комната")
ttk.Combobox(app, textvariable=type_var, values=["Комната", "Квартира", "Многоэтажный дом"], font='Times 14').grid(row=5, column=1)

tk.Button(app, text="Рассчитать", font='Times 14', command=calculate).grid(row=6, column=0, columnspan=2, pady=5)
result_var = tk.StringVar()
tk.Label(app, textvariable=result_var).grid(row=7, column=0, columnspan=2)

tk.Button(app, text="Сохранить в DOC", font='Times 14', command=lambda: save_report("doc")).grid(row=8, column=0, columnspan=2)

app.mainloop()